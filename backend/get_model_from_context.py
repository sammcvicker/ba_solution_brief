from typing import Optional

from pydantic import BaseModel

import io
import mimetypes
from typing import List, Dict, Optional
from pathlib import Path
import logging

# FastAPI imports
from fastapi import UploadFile


# TODO: Group 2 - Define your own BaseModel structure here
# This is a placeholder - create the actual model based on your analysis of the context
class DocumentModel(BaseModel):
    """
    Placeholder model structure - Group 2 should define this based on their analysis.

    Example fields you might consider:
    - title: str
    - summary: str
    - key_points: List[str]
    - categories: List[str]
    - metadata: Dict[str, Any]
    - processed_content: str
    """

    title: Optional[str] = None
    summary: Optional[str] = None
    content: str = ""
    # Add more fields as needed


def get_model_from_context(context: str) -> BaseModel:
    """
    GROUP 2 IMPLEMENTATION:
    Analyze the extracted context and create a structured data model.

    This function should:
    - Parse and analyze the context string from Group 1
    - Extract key information, themes, or data points
    - Structure the information into a well-defined Pydantic model
    - Apply any business logic for categorization or processing
    - Return a populated model instance that Group 3 can use for document generation

    Args:
        context: Raw text content extracted from uploaded documents

    Returns:
        BaseModel: Structured data model containing processed information

    Notes:
    - YOU SHOULD CREATE YOUR OWN MODEL STRUCTURE - the DocumentModel above is just a placeholder
    - Consider what information Group 3 will need to generate the final document
    - Think about data validation and proper typing for your model fields
    - You might want to use AI/LLM services to extract structured data from unstructured text
    - Consider error handling for malformed or insufficient context
    - The model structure will be finalized later, so focus on what makes sense for your processing
    """
    # TODO: Group 2 - Implement context analysis and model creation
    # Create and return your own model instance
    # This is a placeholder implementation
    return DocumentModel(
        title="Processed Document",
        summary="Generated from uploaded files",
        content=context[:100] + "..." if len(context) > 100 else context,
    )


# Document parsing libraries
try:
    import PyPDF2
    from pdfplumber import PDF
except ImportError:
    PyPDF2 = None
    PDF = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import markdown
except ImportError:
    markdown = None

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles extraction of text from various document formats."""

    SUPPORTED_FORMATS = {
        'text/plain': ['txt', 'text', 'log', 'csv'],
        'application/pdf': ['pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['docx'],
        'application/msword': ['doc'],
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['xlsx'],
        'application/vnd.ms-excel': ['xls'],
        'text/markdown': ['md', 'markdown'],
        'text/html': ['html', 'htm'],
        'image/jpeg': ['jpg', 'jpeg'],
        'image/png': ['png'],
        'image/gif': ['gif'],
        'image/bmp': ['bmp']
    }

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        if PDF:
            # Use pdfplumber for better text extraction
            try:
                text_parts = []
                with PDF(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                return '\n\n'.join(text_parts)
            except Exception as e:
                logger.warning(f"pdfplumber failed, falling back to PyPDF2: {e}")

        if PyPDF2:
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_parts.append(page.extract_text())
                return '\n\n'.join(text_parts)
            except Exception as e:
                logger.error(f"Failed to extract text from PDF: {e}")
                return ""

        raise ImportError("PDF processing libraries not installed. Install 'PyPDF2' or 'pdfplumber'")

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file."""
        if not Document:
            raise ImportError("python-docx not installed. Install it with: pip install python-docx")

        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return ""

    @staticmethod
    def extract_text_from_excel(file_content: bytes) -> str:
        """Extract text from Excel file."""
        if not openpyxl:
            raise ImportError("openpyxl not installed. Install it with: pip install openpyxl")

        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_values):
                        text_parts.append(' | '.join(row_values))

            workbook.close()
            return '\n\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text from Excel: {e}")
            return ""

    @staticmethod
    def extract_text_from_image(file_content: bytes) -> str:
        """Extract text from image using OCR."""
        if not Image or not pytesseract:
            raise ImportError("PIL and pytesseract not installed. Install with: pip install pillow pytesseract")

        try:
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from image: {e}")
            return ""

    @staticmethod
    def extract_text_from_html(file_content: bytes) -> str:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(file_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Break into lines and remove leading/trailing space
            lines = (line.strip() for line in text.splitlines())
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            # Drop blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)

            return text
        except ImportError:
            # Fallback to basic extraction
            text = file_content.decode('utf-8', errors='ignore')
            # Remove HTML tags
            import re
            text = re.sub('<[^<]+?>', '', text)
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from HTML: {e}")
            return ""


async def get_context_from_docs(files: List[UploadFile]) -> str:
    """
    Extract and consolidate context/content from uploaded documents.

    This function:
    - Reads and parses various file formats (PDF, DOCX, TXT, etc.)
    - Extracts text content from each file
    - Combines/consolidates the content into a single context string
    - Handles different file types appropriately
    - Returns clean, structured text that can be used for further processing

    Args:
        files: List of uploaded files from the FastAPI endpoint

    Returns:
        str: Consolidated text content from all uploaded files
    """
    if not files:
        return ""

    processor = DocumentProcessor()
    extracted_texts = []

    for file in files:
        try:
            # Read file content
            content = await file.read()
            await file.seek(0)  # Reset file pointer

            # Get file extension
            filename = file.filename or ""
            extension = Path(filename).suffix.lower().lstrip('.')

            # Determine content type
            content_type = file.content_type or mimetypes.guess_type(filename)[0] or 'text/plain'

            logger.info(f"Processing file: {filename} (type: {content_type})")

            # Extract text based on file type
            text = ""

            if extension in ['txt', 'text', 'log', 'csv'] or content_type.startswith('text/plain'):
                # Plain text files
                text = content.decode('utf-8', errors='ignore')

            elif extension == 'pdf' or content_type == 'application/pdf':
                # PDF files
                text = processor.extract_text_from_pdf(content)

            elif extension == 'docx' or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Word documents
                text = processor.extract_text_from_docx(content)

            elif extension in ['xlsx', 'xls'] or content_type in [
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/vnd.ms-excel'
            ]:
                # Excel files
                text = processor.extract_text_from_excel(content)

            elif extension in ['html', 'htm'] or content_type == 'text/html':
                # HTML files
                text = processor.extract_text_from_html(content)

            elif extension in ['md', 'markdown'] or content_type == 'text/markdown':
                # Markdown files
                md_text = content.decode('utf-8', errors='ignore')
                if markdown:
                    text = markdown.markdown(md_text, extensions=['extra'])
                    # Convert HTML to plain text
                    text = processor.extract_text_from_html(text.encode('utf-8'))
                else:
                    text = md_text

            elif extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp'] or content_type.startswith('image/'):
                # Image files (OCR)
                text = processor.extract_text_from_image(content)

            else:
                # Try to decode as text
                try:
                    text = content.decode('utf-8', errors='ignore')
                except:
                    logger.warning(f"Unsupported file type: {filename} ({content_type})")
                    continue

            # Clean and process text
            text = text.strip()

            if text:
                # Add file metadata
                file_header = f"=== File: {filename} ===\n"
                extracted_texts.append(file_header + text)
                logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            else:
                logger.warning(f"No text extracted from {filename}")

        except Exception as e:
            logger.error(f"Error processing file {file.filename}: {str(e)}")
            continue

    # Combine all extracted texts
    if not extracted_texts:
        return "No text content could be extracted from the uploaded documents."

    # Join texts with clear separation
    combined_text = "\n\n" + "=" * 50 + "\n\n".join(extracted_texts) + "\n\n" + "=" * 50

    # Add summary
    summary = f"Successfully processed {len(extracted_texts)} out of {len(files)} files.\n"
    summary += f"Total extracted text length: {len(combined_text)} characters.\n"

    return summary + combined_text


# Additional utility functions
def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    import re

    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)

    # Remove control characters
    text = ''.join(char for char in text if ord(char) >= 32 or char == '\n')

    # Normalize line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def get_text_summary(text: str, max_length: int = 500) -> str:
    """Get a summary/preview of the extracted text."""
    if len(text) <= max_length:
        return text

    # Find a good break point
    break_point = text.rfind(' ', 0, max_length)
    if break_point == -1:
        break_point = max_length

    return text[:break_point] + "..."

